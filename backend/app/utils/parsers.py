"""
Document parsers for PDF and Word files
"""
import PyPDF2
from docx import Document
from typing import List, Tuple
import os


class DocumentParser:
    """Parser for PDF and Word documents"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Tuple[str, int]:
        """
        Parse PDF file and extract text
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        text = ""
        page_count = 0
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                page_count = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num + 1} ---\n"
                        text += page_text
                        
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
        
        return text.strip(), page_count
    
    @staticmethod
    def parse_docx(file_path: str) -> Tuple[str, int]:
        """
        Parse Word document and extract text
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Tuple of (extracted_text, paragraph_count)
        """
        text = ""
        paragraph_count = 0
        
        try:
            doc = Document(file_path)
            paragraph_count = len(doc.paragraphs)
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
                    
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
        
        return text.strip(), paragraph_count
    
    @staticmethod
    def parse_document(file_path: str, file_type: str) -> Tuple[str, int]:
        """
        Parse document based on file type
        
        Args:
            file_path: Path to file
            file_type: File extension ('pdf' or 'docx')
            
        Returns:
            Tuple of (extracted_text, page_or_paragraph_count)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_type.lower() == 'pdf':
            return DocumentParser.parse_pdf(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return DocumentParser.parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
